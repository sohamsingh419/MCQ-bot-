"""PDF source ingestion and retrieval helpers for Source Mode."""
from __future__ import annotations

import asyncio
import hashlib
import re
import shutil
import subprocess
import unicodedata
from pathlib import Path
from typing import Awaitable, Callable, Iterable

from pypdf import PdfReader
from sqlalchemy.exc import IntegrityError
from docx import Document as DocxDocument

from bot.config import Settings, UNIFIED_EXAM_LEVEL, syllabus_topics_for
from bot.database.repositories import Repository
from bot.utils.text import sanitize_text

CHUNK_MAX_CHARS = 4800
MIN_PAGE_TEXT = 40
TEXT_EXTENSIONS = {".txt", ".text", ".md", ".markdown", ".csv", ".tsv", ".json", ".xml", ".html", ".htm", ".rtf", ".log"}
_QUESTION_START_RE = re.compile(r"(?m)^\s*((?:Q(?:uestion)?\s*)?(?:\d{1,4}|[०-९]{1,4})[.)])\s+(.+?)(?=\n\s*(?:\(?[A-Da-d]\)?[.)\-:]|[एबीसीडीए-डी]\s*[.)\-:])\s+)", re.IGNORECASE)
_OPTION_RE = re.compile(r"(?m)^\s*(?:\(?([A-Da-d])\)?|([एबीसीडी]))\s*[.)\-:]\s+(.+?)\s*(?=\n\s*(?:\(?[A-Da-d]\)?|[एबीसीडी])\s*[.)\-:]\s+|\n\s*(?:answer|correct\s+answer|सही\s*उत्तर|उत्तर)\s*[:\-]|\Z)", re.IGNORECASE)
_ANSWER_RE = re.compile(r"(?:answer|ans(?:wer)?|correct\s+answer|सही\s*उत्तर|उत्तर)\s*[:\-]?\s*(?:option\s*)?[\(\[]?\s*([A-Da-dएबीसीडी])", re.IGNORECASE)
ProgressCallback = Callable[[int, int, bool], Awaitable[None]]


def safe_filename(filename: str) -> str:
    name = Path(filename or "source.pdf").name
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip(" .")
    return (name or "source.pdf")[:180]


def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _topic_for_text(text: str, topics: Iterable[str]) -> str:
    # Keep Devanagari combining marks so Hindi OCR can be matched to Hindi topics.
    normalized = re.sub(r"[^\w\s\u0900-\u097F]+", " ", unicodedata.normalize("NFKC", text.casefold()))
    best_topic = "Unclassified"
    best_score = 0
    for topic in topics:
        words = [word for word in re.sub(r"[^\w\s\u0900-\u097F]+", " ", unicodedata.normalize("NFKC", topic.casefold())).split() if len(word) > 3]
        score = sum(1 for word in words if word in normalized)
        if score > best_score:
            best_score = score
            best_topic = topic
    return best_topic


def _ocr_languages() -> set[str]:
    if shutil.which("tesseract") is None:
        return set()
    try:
        result = subprocess.run(
            ["tesseract", "--list-langs"], capture_output=True, text=True, timeout=5, check=False
        )
    except (OSError, subprocess.SubprocessError):
        return set()
    return {
        line.strip() for line in result.stdout.splitlines()
        if line.strip() and not line.lower().startswith("list of available")
    }


def _ocr_page(pdf_path: Path, page_number: int) -> str:
    if shutil.which("pdftoppm") is None or shutil.which("tesseract") is None:
        return ""
    prefix = pdf_path.with_suffix("")
    image_path = Path(f"{prefix}.page-{page_number}")
    try:
        subprocess.run(
            ["pdftoppm", "-f", str(page_number), "-l", str(page_number), "-jpeg", "-scale-to", "1800", str(pdf_path), str(image_path)],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=45,
        )
        generated = Path(f"{image_path}-{'%02d' % page_number}.jpg")
        if not generated.exists():
            candidates = list(pdf_path.parent.glob(f"{image_path.name}-*.jpg"))
            generated = candidates[0] if candidates else generated
        if not generated.exists():
            return ""
        language = "hin+eng" if shutil.which("tesseract") and "hin" in _ocr_languages() else "eng"
        result = subprocess.run(
            ["tesseract", str(generated), "stdout", "-l", language],
            check=True, capture_output=True, text=True, timeout=45,
        )
        generated.unlink(missing_ok=True)
        return sanitize_text(result.stdout).strip()
    except (OSError, subprocess.SubprocessError):
        return ""


def extract_pages(
    pdf_path: Path,
    *,
    ocr_enabled: bool = False,
    progress_callback: Callable[[int, int, bool], None] | None = None,
) -> tuple[int, list[tuple[int, str]]]:
    reader = PdfReader(str(pdf_path))
    total_pages = len(reader.pages)
    pages: list[tuple[int, str]] = []
    for page_number, page in enumerate(reader.pages, 1):
        text = sanitize_text(page.extract_text()).strip()
        used_ocr = False
        if len(text) < MIN_PAGE_TEXT and ocr_enabled:
            used_ocr = True
            text = _ocr_page(pdf_path, page_number).strip()
        if text:
            normalized_page = re.sub(r"[ \t]+", " ", text)
            normalized_page = re.sub(r"\n{3,}", "\n\n", normalized_page).strip()
            pages.append((page_number, normalized_page))
        if progress_callback is not None:
            progress_callback(page_number, total_pages, used_ocr)
    return total_pages, pages


def _decode_text_file(path: Path) -> str:
    raw = path.read_bytes()
    if b"\x00" in raw[:4096]:
        raise ValueError("यह binary file है; readable TXT/CSV/JSON/RTF या DOCX file भेजें।")
    for encoding in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            return sanitize_text(raw.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise ValueError("File का text encoding पढ़ा नहीं जा सका।")


def extract_document_pages(path: Path) -> tuple[int, list[tuple[int, str]]]:
    """Extract readable pages from PDF or common text/document formats."""
    suffix = path.suffix.casefold()
    if suffix == ".pdf":
        return extract_pages(path, ocr_enabled=False)
    if suffix == ".docx":
        document = DocxDocument(str(path))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        text = sanitize_text("\n".join(blocks)).strip()
        return (1, [(1, text)] if text else [])
    if suffix in TEXT_EXTENSIONS or suffix == "":
        text = sanitize_text(_decode_text_file(path)).replace("\r\n", "\n").replace("\r", "\n").strip()
        return (1, [(1, text)] if text else [])
    try:
        text = sanitize_text(_decode_text_file(path)).replace("\r\n", "\n").replace("\r", "\n").strip()
    except ValueError as exc:
        raise ValueError(
            f"{suffix or 'यह'} file readable text/document नहीं है। PDF, TXT, DOCX, CSV, JSON, Markdown या RTF भेजें।"
        ) from exc
    return (1, [(1, text)] if text else [])


async def _extract_document_pages_with_progress(
    path: Path,
    *,
    settings: Settings,
    progress_callback: ProgressCallback | None,
) -> tuple[int, list[tuple[int, str]]]:
    if path.suffix.casefold() == ".pdf":
        return await _extract_pages_with_progress(
            path,
            ocr_enabled=settings.source_ocr_enabled,
            timeout_seconds=getattr(settings, "source_ingest_timeout_seconds", 1800),
            progress_callback=progress_callback,
        )
    pages = await asyncio.to_thread(extract_document_pages, path)
    if progress_callback is not None:
        try:
            total_pages = pages[0]
            await progress_callback(total_pages, total_pages, False)
        except Exception:
            pass
    return pages


def extract_mcq_questions(pages: list[tuple[int, str]]) -> list[dict[str, object]]:
    """Extract already-authored MCQs from text/OCR pages without rewriting them."""
    imported: list[dict[str, object]] = []
    for page_number, page_text in pages:
        text = sanitize_text(page_text).replace("\r", "")
        starts = list(_QUESTION_START_RE.finditer(text))
        for index, match in enumerate(starts):
            block_end = starts[index + 1].start() if index + 1 < len(starts) else len(text)
            block = text[match.start():block_end].strip()
            option_matches = list(_OPTION_RE.finditer(block))
            if len(option_matches) != 4:
                continue
            option_start = option_matches[0].start()
            stem = block[:option_start].strip()
            stem = re.sub(r"^\s*(?:Q(?:uestion)?\s*)?(?:\d{1,4}|[०-९]{1,4})[.)]\s*", "", stem, flags=re.IGNORECASE)
            options = [item.group(3).strip() for item in option_matches]
            answer_match = _ANSWER_RE.search(block[option_matches[-1].end():])
            if not stem or any(not option for option in options) or answer_match is None:
                continue
            answer_token = answer_match.group(1).casefold()
            answer_map = {"a": 0, "b": 1, "c": 2, "d": 3, "ए": 0, "बी": 1, "ब": 1, "सी": 2, "स": 2, "डी": 3, "ड": 3}
            correct_option = answer_map.get(answer_token)
            if correct_option is None:
                continue
            question_type = "Statement-based" if re.search(r"\bstatements?\b|कथनों|कथन\s*[१२३४1-4]", stem, re.IGNORECASE) else "Conceptual"
            language = "Hindi" if re.search(r"[\u0900-\u097F]", f"{stem} {' '.join(options)}") else "English"
            imported.append({
                "question_text": stem,
                "options": options,
                "correct_option": correct_option,
                "question_type": question_type,
                "language": language,
                "page_start": page_number,
                "page_end": page_number,
            })
    return imported


def build_chunks(pages: list[tuple[int, str]], *, state: str, subject: str) -> list[dict[str, object]]:
    topics = syllabus_topics_for(state, subject)
    chunks: list[dict[str, object]] = []
    seen_hashes: set[str] = set()
    current: list[str] = []
    start_page = end_page = 1

    def flush() -> None:
        nonlocal current
        text = sanitize_text(" ".join(current)).strip()
        if text:
            digest = _content_hash(text)
            if digest in seen_hashes:
                current = []
                return
            seen_hashes.add(digest)
            chunks.append({
                "page_start": start_page,
                "page_end": end_page,
                "chunk_index": len(chunks),
                "topic": _topic_for_text(text, topics),
                "text": text,
                "content_hash": digest,
            })
        current = []

    for page_number, page_text in pages:
        paragraphs = [item.strip() for item in re.split(r"(?<=[.!?।])\s+", page_text) if item.strip()]
        for paragraph in paragraphs:
            if not current:
                start_page = page_number
            if current and len(" ".join(current)) + len(paragraph) + 1 > CHUNK_MAX_CHARS:
                end_page = page_number
                flush()
                start_page = page_number
            current.append(paragraph)
            end_page = page_number
    flush()
    return chunks


async def _extract_pages_with_progress(
    pdf_path: Path,
    *,
    ocr_enabled: bool,
    timeout_seconds: int,
    progress_callback: ProgressCallback | None,
) -> tuple[int, list[tuple[int, str]]]:
    loop = asyncio.get_running_loop()
    progress_queue: asyncio.Queue[tuple[int, int, bool]] = asyncio.Queue()

    def emit(page_number: int, total_pages: int, used_ocr: bool) -> None:
        loop.call_soon_threadsafe(progress_queue.put_nowait, (page_number, total_pages, used_ocr))

    worker = asyncio.create_task(asyncio.to_thread(
        extract_pages,
        pdf_path,
        ocr_enabled=ocr_enabled,
        progress_callback=emit,
    ))
    deadline = loop.time() + max(60, timeout_seconds)
    while not worker.done():
        remaining = deadline - loop.time()
        if remaining <= 0:
            raise TimeoutError(f"PDF indexing timeout after {timeout_seconds} seconds")
        try:
            page_number, total_pages, used_ocr = await asyncio.wait_for(
                progress_queue.get(), timeout=min(1.0, remaining)
            )
        except asyncio.TimeoutError:
            continue
        if progress_callback is not None:
            try:
                await progress_callback(page_number, total_pages, used_ocr)
            except Exception:
                pass
    result = await worker
    while not progress_queue.empty():
        page_number, total_pages, used_ocr = progress_queue.get_nowait()
        if progress_callback is not None:
            try:
                await progress_callback(page_number, total_pages, used_ocr)
            except Exception:
                pass
    return result


async def ingest_source_document(
    database,
    settings: Settings,
    document_id: int,
    *,
    state: str,
    subject: str,
    progress_callback: ProgressCallback | None = None,
) -> tuple[bool, str]:
    async with database.session_factory() as session:
        repo = Repository(session)
        document = await repo.get_source_document(document_id)
        if document is None:
            return False, "Source document नहीं मिला।"
        await repo.update_source_document(document_id, state=state, subject=subject, status="processing", extraction_error=None)
        await repo.commit()
        path = Path(document.storage_path)
        try:
            page_count, pages = await _extract_document_pages_with_progress(
                path,
                settings=settings,
                progress_callback=progress_callback,
            )
            combined = "\n".join(text for _, text in pages)
            if len(combined.strip()) < 100:
                raise ValueError("PDF में readable text नहीं मिला। Scanned PDF के लिए SOURCE_OCR_ENABLED=true चाहिए।")
            content_hash = _content_hash(combined)
            existing = await repo.source_document_by_hash(content_hash)
            if existing and existing.id != document_id:
                await repo.update_source_document(document_id, status="duplicate", page_count=page_count, content_hash=content_hash)
                await repo.commit()
                return False, f"यह PDF पहले से source bank में मौजूद है (document {existing.id})।"
            chunks = build_chunks(pages, state=state, subject=subject)
            if not chunks:
                raise ValueError("PDF से usable text chunks नहीं बने।")
            count = await repo.replace_source_chunks(document_id, state=state, subject=subject, chunks=chunks)
            imported_questions = 0
            skipped_duplicates = 0
            for item in extract_mcq_questions(pages):
                answer_label = chr(ord("A") + int(item["correct_option"]))
                topic = _topic_for_text(str(item["question_text"]), syllabus_topics_for(state, subject))
                try:
                    async with session.begin_nested():
                        await repo.add_question(
                            question_text=str(item["question_text"]), options=list(item["options"]),
                            correct_option=int(item["correct_option"]),
                            explanation=f"PDF answer key: option {answer_label}",
                            key_point=f"PDF answer key: option {answer_label}",
                            state=state, subject=subject, topic=topic, difficulty=UNIFIED_EXAM_LEVEL,
                            question_type=str(item["question_type"]), language=str(item["language"]),
                            source="source", source_group_id=document.telegram_chat_id,
                            source_document_id=document_id, source_title=document.filename,
                            source_page_start=int(item["page_start"]), source_page_end=int(item["page_end"]),
                        )
                    imported_questions += 1
                except IntegrityError:
                    skipped_duplicates += 1
            await repo.update_source_document(
                document_id, status="ready", page_count=page_count, chunk_count=count,
                content_hash=content_hash, title=document.filename,
            )
            await repo.commit()
            return True, (
                f"Source तैयार है: {document.filename} • {count} sections • {page_count} pages • "
                f"{imported_questions} MCQs imported verbatim"
                + (f" • {skipped_duplicates} duplicates skipped" if skipped_duplicates else "")
            )
        except Exception as exc:
            await repo.update_source_document(document_id, status="failed", extraction_error=str(exc)[:1000])
            await repo.commit()
            return False, f"PDF process नहीं हो सकी: {exc}"
