from pathlib import Path

import pytest
from docx import Document

from bot.services.source import extract_document_pages


def test_extract_txt_source_file(tmp_path: Path) -> None:
    path = tmp_path / "rajasthan_history.txt"
    path.write_text("राजस्थान का इतिहास\n\n1. पहला तथ्य।", encoding="utf-8")
    page_count, pages = extract_document_pages(path)
    assert page_count == 1
    assert pages == [(1, "राजस्थान का इतिहास\n\n1. पहला तथ्य।")]


def test_extract_docx_source_file_including_table(tmp_path: Path) -> None:
    path = tmp_path / "questions.docx"
    document = Document()
    document.add_paragraph("राजस्थान का राज्य पक्षी कौन सा है?")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A. मोर"
    table.rows[0].cells[1].text = "B. गोडावण"
    document.save(path)
    page_count, pages = extract_document_pages(path)
    assert page_count == 1
    assert "राजस्थान का राज्य पक्षी कौन सा है?" in pages[0][1]
    assert "A. मोर | B. गोडावण" in pages[0][1]


def test_extract_rejects_unknown_binary_document(tmp_path: Path) -> None:
    path = tmp_path / "archive.zip"
    path.write_bytes(b"PK\x03\x04\x00\x00")
    with pytest.raises(ValueError, match="readable text/document"):
        extract_document_pages(path)
